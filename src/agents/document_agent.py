"""
Agente principal de procesamiento de documentos
"""
import asyncio
import time
from typing import Dict, List, Optional
from loguru import logger
from pathlib import Path

from src.models.schemas import (
    Document, ProcessingResult, AgentRequest, AgentResponse,
    ProcessingStatus, DocumentType, ClassificationResult,
    ExtractionResult, ValidationResult, FraudDetectionResult
)
from src.services.ocr_service import OCRService
from src.services.classification_service import DocumentClassifier
from src.services.extraction_service import DataExtractionService
from src.core.config import DOCUMENT_CONFIG


class DocumentProcessingAgent:
    """Agente principal para procesamiento inteligente de documentos"""
    
    def __init__(self):
        self.ocr_service = OCRService()
        self.classifier = DocumentClassifier()
        self.extractor = DataExtractionService()
        
        # Cargar modelo de clasificación si existe
        self.classifier.load_model()
        
        logger.info("Agente de procesamiento de documentos inicializado")
    
    async def process_document(self, document: Document, actions: List[str] = None) -> ProcessingResult:
        """
        Procesa un documento completo
        
        Args:
            document: Documento a procesar
            actions: Lista de acciones a realizar ['classify', 'extract', 'validate', 'detect_fraud']
        """
        start_time = time.time()
        
        if actions is None:
            actions = ['classify', 'extract', 'validate']
        
        result = ProcessingResult(
            document=document,
            errors=[]
        )
        
        try:
            # Actualizar estado
            document.status = ProcessingStatus.PROCESSING
            
            # 1. Extraer texto del documento
            text = await self._extract_text(document)
            if not text:
                result.errors.append("No se pudo extraer texto del documento")
                document.status = ProcessingStatus.FAILED
                return result
            
            # 2. Clasificar documento
            if 'classify' in actions:
                logger.info(f"Clasificando documento {document.id}")
                result.classification = self._classify_document(text)
                document.document_type = result.classification.document_type
            
            # 3. Extraer datos
            if 'extract' in actions and document.document_type:
                logger.info(f"Extrayendo datos del documento {document.id}")
                result.extraction = self._extract_data(text, document.document_type)
            
            # 4. Validar documento
            if 'validate' in actions:
                logger.info(f"Validando documento {document.id}")
                result.validation = self._validate_document(result.extraction, document.document_type)
            
            # 5. Detectar fraudes
            if 'detect_fraud' in actions:
                logger.info(f"Analizando fraudes en documento {document.id}")
                result.fraud_detection = self._detect_fraud(text, result.extraction)
            
            # Actualizar estado final
            if result.errors:
                document.status = ProcessingStatus.FAILED
            elif result.validation and not result.validation.is_valid:
                document.status = ProcessingStatus.REJECTED
            else:
                document.status = ProcessingStatus.COMPLETED
            
            result.processing_time = time.time() - start_time
            
            logger.info(f"Documento {document.id} procesado en {result.processing_time:.2f}s")
            
        except Exception as e:
            logger.error(f"Error procesando documento {document.id}: {e}")
            result.errors.append(f"Error en procesamiento: {str(e)}")
            document.status = ProcessingStatus.FAILED
        
        return result
    
    async def _extract_text(self, document: Document) -> Optional[str]:
        """Extrae texto del documento según su tipo"""
        try:
            file_extension = Path(document.file_path).suffix.lower()
            
            if file_extension == '.pdf':
                # Extraer de PDF
                ocr_result = self.ocr_service.extract_from_pdf(document.file_path)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
                # Extraer de imagen
                ocr_result = self.ocr_service.extract_text(document.file_path)
            elif file_extension == '.docx':
                # Extraer de Word
                from docx import Document as DocxDocument
                doc = DocxDocument(document.file_path)
                text = '\\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text
            else:
                logger.warning(f"Tipo de archivo no soportado: {file_extension}")
                return None
            
            if ocr_result.get('confidence', 0) < 0.3:
                logger.warning(f"Baja confianza en OCR: {ocr_result.get('confidence', 0)}")
            
            return ocr_result.get('text', '')
            
        except Exception as e:
            logger.error(f"Error extrayendo texto: {e}")
            return None
    
    def _classify_document(self, text: str) -> ClassificationResult:
        """Clasifica el tipo de documento"""
        return self.classifier.classify(text)
    
    def _extract_data(self, text: str, document_type: DocumentType) -> ExtractionResult:
        """Extrae datos estructurados del documento"""
        return self.extractor.extract_data(text, document_type)
    
    def _validate_document(self, extraction: Optional[ExtractionResult], 
                          document_type: DocumentType) -> ValidationResult:
        """Valida la información extraída del documento"""
        errors = []
        warnings = []
        compliance_score = 1.0
        
        if not extraction or not extraction.fields:
            errors.append("No se extrajo información del documento")
            compliance_score = 0.0
        else:
            # Validaciones específicas por tipo de documento
            if document_type == DocumentType.CEDULA:
                if 'numero_documento' not in extraction.fields:
                    errors.append("Número de documento no encontrado")
                    compliance_score -= 0.3
                elif len(extraction.fields['numero_documento']) < 6:
                    errors.append("Número de documento inválido")
                    compliance_score -= 0.2
                
                if 'nombres' not in extraction.fields:
                    warnings.append("Nombres no extraídos claramente")
                    compliance_score -= 0.1
            
            elif document_type == DocumentType.ESTADO_CUENTA:
                if 'numero_cuenta' not in extraction.fields:
                    errors.append("Número de cuenta no encontrado")
                    compliance_score -= 0.4
                
                if 'saldo' not in extraction.fields:
                    warnings.append("Saldo no identificado")
                    compliance_score -= 0.1
            
            # Validar confianza de extracción
            avg_confidence = sum(extraction.confidence_scores.values()) / len(extraction.confidence_scores) if extraction.confidence_scores else 0
            if avg_confidence < 0.5:
                warnings.append(f"Baja confianza en extracción: {avg_confidence:.2f}")
                compliance_score -= 0.2
        
        compliance_score = max(0.0, compliance_score)
        
        return ValidationResult(
            is_valid=len(errors) == 0,
            errors=errors,
            warnings=warnings,
            compliance_score=compliance_score
        )
    
    def _detect_fraud(self, text: str, extraction: Optional[ExtractionResult]) -> FraudDetectionResult:
        """Detecta posibles fraudes en el documento"""
        risk_factors = []
        risk_score = 0.0
        recommendations = []
        
        # Análisis básico de fraude
        text_lower = text.lower()
        
        # Verificar calidad del OCR
        if extraction and extraction.confidence_scores:
            avg_confidence = sum(extraction.confidence_scores.values()) / len(extraction.confidence_scores)
            if avg_confidence < 0.4:
                risk_factors.append("Baja calidad de imagen/documento")
                risk_score += 0.3
        
        # Verificar patrones sospechosos
        suspicious_patterns = [
            "photoshop", "editado", "modificado", "alterado",
            "copia", "duplicado", "falso", "fraudulento"
        ]
        
        for pattern in suspicious_patterns:
            if pattern in text_lower:
                risk_factors.append(f"Texto sospechoso detectado: {pattern}")
                risk_score += 0.4
        
        # Verificar inconsistencias en fechas
        if extraction and extraction.fields:
            dates = [v for k, v in extraction.fields.items() if 'fecha' in k.lower()]
            # Aquí se podrían implementar validaciones de fechas más complejas
        
        # Verificar formato de números de documento
        if extraction and 'numero_documento' in extraction.fields:
            doc_number = extraction.fields['numero_documento']
            if not doc_number.isdigit() or len(doc_number) < 6:
                risk_factors.append("Formato inválido de número de documento")
                risk_score += 0.2
        
        # Generar recomendaciones
        if risk_score > 0.5:
            recommendations.append("Revisar manualmente el documento")
            recommendations.append("Verificar autenticidad con fuentes oficiales")
        
        if risk_score > 0.3:
            recommendations.append("Solicitar documentos adicionales")
        
        risk_score = min(1.0, risk_score)
        
        return FraudDetectionResult(
            is_fraudulent=risk_score > 0.6,
            risk_score=risk_score,
            risk_factors=risk_factors,
            recommendations=recommendations
        )
    
    async def process_request(self, request: AgentRequest) -> AgentResponse:
        """Procesa una solicitud completa"""
        try:
            # Aquí normalmente cargarías el documento desde una base de datos
            # Por ahora, creamos un documento de ejemplo
            document = Document(
                id=request.document_id,
                filename=f"document_{request.document_id}",
                file_path=f"/tmp/{request.document_id}",
                file_size=1024,
                mime_type="application/pdf"
            )
            
            result = await self.process_document(document, request.actions)
            
            return AgentResponse(
                request_id=request.document_id,
                status=document.status,
                result=result,
                message="Procesamiento completado exitosamente"
            )
            
        except Exception as e:
            logger.error(f"Error procesando solicitud {request.document_id}: {e}")
            return AgentResponse(
                request_id=request.document_id,
                status=ProcessingStatus.FAILED,
                result=None,
                message=f"Error en procesamiento: {str(e)}"
            )
